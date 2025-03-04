#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Modul: tab_agent_convo.py
Dieses Modul enthält die Klasse AgentConversationTab, die den UI-Block für die Agenten-Konversation
sowohl als eigenständige Gradio-Anwendung als auch als Tab bereitstellt.
Entwickelt von CipherCore.
"""

import gradio as gr
import re
import logging
import datetime
import json
import hashlib
import io
import os
import time
from collections import defaultdict
from typing import List, Dict, Tuple, Any
import sqlite3
from jsonschema import validate, ValidationError
from docx import Document
from docx.shared import Inches

from dotenv import load_dotenv
from google import genai


class AgentConversationTab:
    def __init__(self):
        """Initialisiert den AgentConversationTab."""
        load_dotenv()

        self.api_key = os.getenv("API_KEY")
        self.model_name = "gemini-2.0-flash-thinking-exp-01-21"

        if not self.api_key:
            raise ValueError("API_KEY nicht in .env gefunden!")

        self.client = genai.Client(api_key=self.api_key)

        logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

        self.api_sleep_seconds = 10
        self.api_max_retries = 3

        self.audit_log_file = "audit_log.txt"
        self.expiration_time_seconds = 300
        self.role_permissions = {
            "user": ["REQ", "DATA"],
            "admin": ["REQ", "DATA", "CALC", "IF", "AI"]
        }
        self.priority_map = {"HIGH": 1, "MEDIUM": 2, "LOW": 3}

        self.user_data_file = "user_data.json"
        self.discussion_db_file = "discussion_data.db"
        self.rating_data_file = "rating_data.json"
        self.agent_config_file = "agent_config.json"

        self.user_data_schema = {
            "type": "object",
            "patternProperties": {
                "^[a-zA-Z0-9_-]+$": {
                    "type": "object",
                    "properties": {
                        "password": {"type": "string"}
                    },
                    "required": ["password"]
                }
            }
        }

        self.agent_config_schema = {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "name": {"type": "string"},
                    "personality": {"type": "string", "enum": ["kritisch", "visionär", "konservativ", "neutral"]},
                    "description": {"type": "string"}
                },
                "required": ["name", "personality", "description"]
            }
        }
        self.discussion_ratings = defaultdict(lambda: defaultdict(dict), self.load_rating_data())


    # --- Hilfsfunktionen ---
    def load_json_data(self, filename: str, schema: dict = None) -> Dict[str, Any]:
        try:
            with open(filename, "r", encoding="utf-8") as f:
                data = json.load(f)
                if schema:
                    validate(instance=data, schema=schema)
                return data
        except FileNotFoundError:
            logging.warning(f"Datei '{filename}' nicht gefunden. Starte mit leeren Daten.")
            return {}
        except json.JSONDecodeError as e:
            logging.error(f"Fehler beim Lesen von '{filename}': Ungültiges JSON-Format. Details: {e}")
            return {}
        except ValidationError as e:
            logging.error(f"Datei '{filename}' entspricht nicht dem erwarteten Schema: {e}")
            return {}

    def save_json_data(self, data: Dict[str, Any], filename: str) -> None:
        try:
            with open(filename, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=4)
        except IOError as e:
            logging.error(f"Fehler beim Schreiben in Datei '{filename}': {e}")

    def load_user_data(self) -> Dict[str, Any]:
        return self.load_json_data(self.user_data_file, self.user_data_schema)

    def save_user_data(self, user_data: Dict[str, Any]) -> None:
        self.save_json_data(user_data, self.user_data_file)

    def load_rating_data(self) -> Dict[str, Any]:
        return self.load_json_data(self.rating_data_file)

    def save_rating_data(self, rating_data: Dict[str, Any]) -> None:
        self.save_json_data(rating_data, self.rating_data_file)

    def load_agent_config(self) -> List[Dict[str, str]]:
        config = self.load_json_data(self.agent_config_file, self.agent_config_schema)
        if not isinstance(config, list):
            logging.error(f"Agentenkonfiguration in '{self.agent_config_file}' ist ungültig oder leer. Stelle sicher, dass es eine Liste ist.")
            return []
        return config

    def hash_password(self, password: str) -> str:
        return hashlib.sha256(password.encode()).hexdigest()

    def verify_password(self, password: str, hashed_password: str) -> bool:
        return self.hash_password(password) == hashed_password

    def register_user(self, username: str, password: str) -> str:
        if not re.match(r"^[a-zA-Z0-9_-]+$", username):
            return "Ungültiger Nutzername. Nur Buchstaben, Zahlen, '-', '_' erlaubt."
        if len(password) < 8:
            return "Passwort muss mindestens 8 Zeichen lang sein."

        user_data = self.load_user_data()
        if username in user_data:
            return "Nutzername bereits vergeben."
        user_data[username] = {"password": self.hash_password(password)}
        self.save_user_data(user_data)
        return "Registrierung erfolgreich."

    def login_user(self, username: str, password: str) -> Tuple[str, str]:
        """ Gibt (Meldung, username_oder_None) zurück """
        user_data = self.load_user_data()
        if username in user_data and self.verify_password(password, user_data[username]["password"]):
            return "Login erfolgreich.", username
        return "Login fehlgeschlagen.", None

    def create_discussion_table(self):
        """ Erstellt die Diskussionstabelle, falls nicht vorhanden. """
        conn = sqlite3.connect(self.discussion_db_file)
        cursor = conn.cursor()
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS discussions (
                discussion_id TEXT PRIMARY KEY,
                topic TEXT,
                agents TEXT,
                chat_history TEXT,
                summary TEXT,
                user TEXT,
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
            )
        """)
        conn.commit()
        conn.close()

    def save_discussion_data_db(self, discussion_id: str, topic: str, agents: List[str], chat_history: List[Dict], summary: str, user: str = None) -> None:
        """ Speichert Diskussionsdaten in der SQLite Datenbank. """
        conn = sqlite3.connect(self.discussion_db_file)
        cursor = conn.cursor()
        try:
            cursor.execute("""
                INSERT INTO discussions (discussion_id, topic, agents, chat_history, summary, user)
                VALUES (?, ?, ?, ?, ?, ?)
            """, (discussion_id, topic, json.dumps(agents), json.dumps(chat_history), summary, user))
            conn.commit()
        except sqlite3.Error as e:
            logging.error(f"Datenbankfehler beim Speichern der Diskussion '{discussion_id}': {e}")
            conn.rollback()
        finally:
            conn.close()

    def load_discussion_data_db(self, user: str = None) -> Dict[str, Any]:
        """ Lädt Diskussionsdaten aus der SQLite Datenbank. Optional für einen bestimmten Nutzer. """
        conn = sqlite3.connect(self.discussion_db_file)
        cursor = conn.cursor()
        discussions = {}
        try:
            if user:
                cursor.execute("SELECT * FROM discussions WHERE user = ?", (user,))
            else:
                cursor.execute("SELECT * FROM discussions")
            rows = cursor.fetchall()
            for row in rows:
                disc_id, topic, agents_json, chat_history_json, summary, user_name, timestamp = row
                agents = json.loads(agents_json) if agents_json else []
                chat_history = json.loads(chat_history_json) if chat_history_json else []
                discussions[disc_id] = {
                    "topic": topic,
                    "agents": agents,
                    "chat_history": chat_history,
                    "summary": summary,
                    "user": user_name,
                    "timestamp": timestamp
                }
        except sqlite3.Error as e:
            logging.error(f"Datenbankfehler beim Laden der Diskussionen: {e}")
        finally:
            conn.close()
        return discussions

    def evaluate_response(self, response: str) -> str:
        resp_l = response.lower()
        if "wiederhole mich" in resp_l:
            return "schlechte antwort"
        elif "neue perspektive" in resp_l:
            return "gute antwort"
        else:
            return "neutral"

    def rate_agent_response(self, discussion_id: str, iteration: int, agent_name: str, rating_type: str) -> None:
        """
        Verhindert KeyError, indem wir (discussion_id, iteration, agent_name) bei Bedarf anlegen.
        """
        if agent_name not in self.discussion_ratings[discussion_id][iteration]:
            self.discussion_ratings[discussion_id][iteration][agent_name] = {"upvotes": 0, "downvotes": 0}

        if rating_type == "upvote":
            self.discussion_ratings[discussion_id][iteration][agent_name]["upvotes"] += 1
        elif rating_type == "downvote":
            self.discussion_ratings[discussion_id][iteration][agent_name]["downvotes"] += 1

        self.save_rating_data(self.discussion_ratings)

    def call_gemini_api(self, prompt: str) -> Dict[str, str]:
        """
        Ruft die Gemini-API auf mit erweitertem Retry-Mechanismus und Fehlerbehandlung.
        """
        retry_delay = self.api_sleep_seconds
        for attempt in range(self.api_max_retries + 1):
            try:
                logging.info(f"[{attempt+1}/{self.api_max_retries+1}] Sende Prompt an Gemini: {prompt[:100]}...")
                response = self.client.models.generate_content(model=self.model_name, contents=[prompt])

                # Wartezeit nach jedem Request
                time.sleep(self.api_sleep_seconds)

                if not hasattr(response, "text") or not response.text:
                    msg = "Leere Antwort von Gemini API."
                    logging.warning(msg)
                    return {"response": msg}

                return {"response": response.text}

            except genai.APIError as e:
                err_s = str(e)
                logging.error(f"Gemini API Fehler (Versuch {attempt+1}): {err_s}, Status Code: {e.status_code}")

                if e.status_code == 429:
                    if attempt < self.api_max_retries:
                        retry_delay *= 2
                        logging.info(f"Rate Limit erreicht. Warte {retry_delay}s und versuche erneut.")
                        time.sleep(retry_delay)
                        continue
                    else:
                        return {"response": f"API Rate Limit erreicht nach mehreren Versuchen. Bitte später erneut versuchen."}
                elif e.status_code == 503:
                    if attempt < self.api_max_retries:
                        logging.info(f"Server überlastet. Warte {retry_delay}s und versuche erneut.")
                        time.sleep(retry_delay)
                        continue
                    else:
                        return {"response": "Gemini API Server überlastet nach mehreren Versuchen."}
                elif e.status_code == 401:
                    return {"response": "Authentifizierungsfehler bei der Gemini API. Überprüfen Sie den API-Schlüssel."}
                else:
                    if attempt < self.api_max_retries:
                        logging.info(f"Unerwarteter API Fehler, versuche Retry. Warte {retry_delay}s.")
                        time.sleep(retry_delay)
                        continue
                    else:
                        return {"response": f"Unerwarteter Fehler bei Gemini API Aufruf nach mehreren Versuchen: {err_s}"}
            except Exception as e:
                err_s = str(e)
                logging.error(f"Genereller Fehler bei Gemini API Aufruf (Versuch {attempt+1}): {err_s}")
                return {"response": f"Fehler bei Gemini API Aufruf: {err_s}"}

        return {"response": "Unbekannter Fehler nach mehreren API-Versuchen."}

    def joint_conversation_with_selected_agents(
        self,
        conversation_topic: str,
        selected_agents: List[Dict[str, str]],
        iterations: int,
        expertise_level: str,
        language: str,
        chat_history: List[Dict[str, str]],
        user_state: str,
        discussion_id: str = None
    ):
        """
        user_state: der eingeloggte Nutzername (oder None).
        """
        if discussion_id is None:
            discussion_id = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")

        active_agents_names = [sa["name"] for sa in selected_agents]
        num_agents = len(active_agents_names)
        agent_outputs = [""] * num_agents
        topic_changed = False

        logging.info(f"Konversation gestartet: {active_agents_names}, iters={iterations}, level={expertise_level}, lang={language}, Diskussions-ID: {discussion_id}")

        for i in range(iterations):
            agent_idx = i % num_agents
            current_agent_name = active_agents_names[agent_idx]
            current_agent_config = next((a for a in selected_agents if a["name"] == current_agent_name), None)
            current_personality = current_agent_config.get("personality", "neutral")
            current_instruction = current_agent_config.get("instruction", "")

            prev_agent_name = active_agents_names[(agent_idx - 1) % num_agents]
            prev_output = agent_outputs[(agent_idx - 1) % num_agents]

            prompt_txt = (
                f"Wir führen eine Konversation über: '{conversation_topic}'.\n"
                f"Iteration {i+1}: Agent {current_agent_name} (Spezialist für **{current_agent_name}**). {current_instruction}\n"
                f"Agent {prev_agent_name} sagte: {prev_output}\n"
            )

            if current_personality == "kritisch":
                prompt_txt += "\nSei kritisch und hinterfrage Annahmen."
            elif current_personality == "visionär":
                prompt_txt += "\nSei visionär und denke groß."
            elif current_personality == "konservativ":
                prompt_txt += "\nSei konservativ und bleibe bei Bewährtem."

            if language == "Deutsch":
                prompt_txt += "\n\nAntworte auf Deutsch."
            elif language == "Englisch":
                prompt_txt += "\n\nRespond in English."
            elif language == "Französisch":
                prompt_txt += "\n\nRépondez en français."
            elif language == "Spanisch":
                prompt_txt += "\n\nResponde en español."

            chat_history.append({
                "role": "user",
                "content": f"Agent {current_agent_name} (Iteration {i+1}): Thema {conversation_topic}, vorheriger: {prev_agent_name}: {prev_output}"
            })

            api_resp = self.call_gemini_api(prompt_txt)
            agent_output = api_resp.get("response", f"Keine Antwort von {current_agent_name}")
            agent_outputs[agent_idx] = agent_output

            qual = self.evaluate_response(agent_output)
            if qual == "schlechte antwort":
                logging.info(f"{current_agent_name} => 'schlechte antwort', retry ...")
                retry_resp = self.call_gemini_api("Versuche eine kreativere Antwort.")
                retry_output = retry_resp.get("response", f"Keine Retry-Antwort von {current_agent_name}")
                if "Fehler bei Gemini API Aufruf" not in retry_output:
                    agent_output = retry_output
                agent_outputs[agent_idx] = agent_output

            chat_history.append({
                "role": "assistant",
                "content": f"Antwort von Agent {current_agent_name} (Iteration {i+1}):\n{agent_output}"
            })
            logging.info(f"Antwort Agent {current_agent_name} (i={i+1}): {agent_output}")

            formatted_output_chunk = (
                f"**Iteration {i+1}: Agent {current_agent_name} ({current_personality})**\n\n"
                f"{agent_output}\n\n"
                "---\n\n"
            )

            yield chat_history, formatted_output_chunk, discussion_id, (i+1), current_agent_name

            if i > (iterations * 0.6) and agent_output == agent_outputs[(agent_idx - 1) % num_agents] and not topic_changed:
                new_topic = "Neues Thema: KI-Trends 2026"
                agent_outputs = [new_topic] * num_agents
                topic_changed = True

        sum_prompt = f"Fasse die gesamte Diskussion über '{conversation_topic}' zusammen."
        sum_resp = self.call_gemini_api(sum_prompt)
        sum_text = sum_resp.get("response", "Keine Zusammenfassung generiert.")
        chat_history.append({
            "role": "assistant",
            "content": f"**Zusammenfassung**:\n{sum_text}"
        })

        if user_state:
            self.save_discussion_data_db(discussion_id, conversation_topic, active_agents_names, chat_history, sum_text, user_state)
            logging.info(f"Diskussion {discussion_id} für {user_state} in Datenbank gespeichert.")
        else:
            logging.info("Keine Speicherung in Datenbank, kein Benutzer eingeloggt.")

        final_text = agent_outputs[-1]
        chat_history.append({
            "role": "assistant",
            "content": f"Finale Aussage:\n{final_text}"
        })

        logging.info(f"Finale Aussage: {final_text}")

        yield chat_history, sum_text, discussion_id, None, None

    def save_chat_as_word(self, chat_history: List[Dict], discussion_id: str) -> str:
        """Speichert den Chatverlauf als formatierte Word-Datei."""
        document = Document()
        document.add_heading(f'CipherCore Agenten-Diskussion: {discussion_id}', level=1)

        for message in chat_history:
            role = message['role']
            content = message['content']
            if role == 'user':
                document.add_paragraph(f"Nutzer:", style='List Bullet').add_run(f" {content}").bold = True
            elif role == 'assistant':
                agent_name_match = re.search(r'Agent (.*?)\s', content) # Agentennamen extrahieren
                agent_name = agent_name_match.group(1) if agent_name_match else "Agent"
                p = document.add_paragraph(f"{agent_name}:", style='List Bullet')
                p.add_run(f" {content.split(':\n', 1)[1] if ':\n' in content else content}") # Antwortinhalt extrahieren

        filename = f"CipherCore_Diskussion_{discussion_id}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        try:
            document.save(filename)
            logging.info(f"Word-Datei '{filename}' erfolgreich gespeichert.")
            return filename
        except Exception as e:
            logging.error(f"Fehler beim Speichern der Word-Datei: {e}")
            return None # Fehlerfall signalisieren


    def build_tab(self):
        """Baut den UI-Block für den Agenten-Konversationstab."""
        self.create_discussion_table() # Stelle sicher, dass die Tabelle existiert

        with gr.TabItem("Agenten-Konversation"):
            gr.Markdown("## Willkommen bei CipherCore! Ihre Plattform für sichere Programmierung und innovative KI-Lösungen.")
            gr.Markdown("### Agenten-Konversation")
            gr.Markdown("Dieses Tool demonstriert eine Konversation zwischen verschiedenen KI-Agenten, die von CipherCore für Sie entwickelt wurden. Wählen Sie Agenten aus, geben Sie ein Thema vor und starten Sie die Diskussion. Wir bei CipherCore legen größten Wert auf Sicherheit und Innovation in allen unseren Lösungen.")

            # State für user_name
            user_state = gr.State(value=None)

            with gr.Row():
                with gr.Column():
                    gr.Markdown("### Login")
                    username_login = gr.Textbox(label="Nutzername")
                    password_login = gr.Textbox(label="Passwort", type="password")
                    login_btn = gr.Button("Login")
                    login_status_label = gr.Label(label="Login-Status")

                with gr.Column():
                    gr.Markdown("### Registrierung")
                    username_register = gr.Textbox(label="Nutzername")
                    password_register = gr.Textbox(label="Passwort", type="password")
                    register_btn = gr.Button("Registrieren")
                    register_status_label = gr.Label(label="Registrierungs-Status")

            def login_event(u_name, u_pass, current_usr):
                msg, logged_in_user = self.login_user(u_name, u_pass)
                if logged_in_user:
                    return msg, logged_in_user
                else:
                    return msg, current_usr

            def register_event(u_name, u_pass):
                return self.register_user(u_name, u_pass)

            login_btn.click(
                fn=login_event,
                inputs=[username_login, password_login, user_state],
                outputs=[login_status_label, user_state]
            )

            register_btn.click(
                fn=register_event,
                inputs=[username_register, password_register],
                outputs=[register_status_label]
            )

            gr.Markdown("---")

            # Agenten-Liste DYNAMISCH LADEN
            agent_config_data = self.load_agent_config()
            agent_checkboxes = []
            agent_personalities_radios = []

            with gr.Accordion("Agenten Auswahl (auf-/zuklappbar)", open=False):
                gr.Markdown("### Wähle Agenten und deren Persönlichkeit:")
                with gr.Column():
                    for agent_data in agent_config_data:
                        with gr.Row():
                            cbox = gr.Checkbox(label=agent_data["name"])
                            radio = gr.Radio(["kritisch", "visionär", "konservativ", "neutral"], value="kritisch", label="Persönlichkeit")
                            agent_checkboxes.append(cbox)
                            agent_personalities_radios.append(radio)

            iteration_slider = gr.Slider(20, 100, value=10, step=1, label="Anzahl Gesprächsrunden")
            level_radio = gr.Radio(["Beginner", "Fortgeschritten", "Experte"], value="Experte", label="Experten-Level")
            lang_radio = gr.Radio(["Deutsch", "Englisch", "Französisch", "Spanisch"], value="Deutsch", label="Sprache")

            topic_input = gr.Textbox(label="Diskussionsthema")

            # Gespeicherte Diskussionen ansehen
            with gr.Accordion("Gespeicherte Diskussionen", open=False):
                saved_disc_label = gr.JSON()
                load_disc_btn = gr.Button("Diskussionen laden")

                def load_discs(current_usr):
                    if current_usr:
                        disc_data = self.load_discussion_data_db(current_usr)
                    else:
                        disc_data = {"Warnung": "Niemand eingeloggt."}
                    return disc_data


                load_disc_btn.click(
                    fn=load_discs,
                    inputs=[user_state],
                    outputs=[saved_disc_label]
                )

            # Chatbot
            chatbot = gr.Chatbot(label="Agenten-Konversation", type="messages")
            formatted_output_md = gr.Markdown(label="Formatierter Output")

            rating_row = gr.Row(visible=False)
            with rating_row:
                upvote_btn = gr.Button("👍")
                downvote_btn = gr.Button("👎")
                rating_label = gr.Label("Bewertung: Nicht bewertet")
                rating_info = gr.State({})

            with gr.Row():
                start_btn = gr.Button("Konversation starten")
                save_btn = gr.Button("Diskussion speichern")
                word_save_btn = gr.Button("Chat als Word speichern") # NEU: Word-Speichern Button

            word_file_output = gr.File(label="Word-Datei herunterladen", visible=False) # Komponente für Word-Datei-Download

            def process_input_custom_agents(
                user_topic, chat_history, rating_state, current_usr_state,
                *args
            ):
                """
                Diese Funktion ruft joint_conversation_with_selected_agents(...) auf
                und yieldet Zwischenergebnisse an Gradio.
                """
                num_agents_config = len(agent_config_data)
                cbox_vals = args[:num_agents_config]
                pers_vals = args[num_agents_config:2*num_agents_config]
                iters = args[2*num_agents_config]
                level = args[2*num_agents_config + 1]
                lang = args[2*num_agents_config + 2]
                discussion_id_state = args[2*num_agents_config + 3]

                selected_agents = []
                for i, agent_data in enumerate(agent_config_data):
                    if cbox_vals[i]:
                        selected_agents.append({
                            "name": agent_data["name"],
                            "personality": pers_vals[i],
                            "instruction": agent_data.get("description", "")
                        })

                if not selected_agents:
                    yield ("Keine Agenten ausgewählt!", "", discussion_id_state, gr.update(visible=False), rating_state)
                    return

                if discussion_id_state is None:
                    discussion_id_state = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")

                agent_convo = self.joint_conversation_with_selected_agents(
                    conversation_topic=user_topic,
                    selected_agents=selected_agents,
                    iterations=iters,
                    expertise_level=level,
                    language=lang,
                    chat_history=chat_history,
                    user_state=current_usr_state,
                    discussion_id=discussion_id_state
                )

                formatted_output_text = ""
                for updated_hist, chunk_text, disc_id, iteration_num, agent_n in agent_convo:
                    formatted_output_text += chunk_text
                    rating_state["discussion_id"] = disc_id
                    rating_state["iteration"] = iteration_num
                    rating_state["agent_name"] = agent_n

                    yield (
                        updated_hist,
                        formatted_output_text,
                        disc_id,
                        gr.update(visible=True),
                        rating_state,
                        gr.update(visible=False) # Word-Datei Output unsichtbar machen während Konversation
                    )

            def save_discussion_manually(
                history, user_topic, dummy_state, fmt_output, discussion_id_in, user_state_in
            ):
                """ Speichert die Diskussion. """
                if user_state_in:
                    active_agents_names = [agent['name'] for agent in self.load_agent_config()]
                    self.save_discussion_data_db(discussion_id_in, user_topic, active_agents_names, history, "Manuell gespeichert", user_state_in)
                    return "Diskussion in Datenbank gespeichert."
                else:
                    return "Bitte zuerst einloggen."

            def save_chat_word_event(history, discussion_id_in):
                """ Event-Handler für den "Chat als Word speichern" Button. """
                if discussion_id_in:
                    word_filename = self.save_chat_as_word(history, discussion_id_in)
                    if word_filename:
                        return gr.File(value=word_filename, label="Word-Datei herunterladen", visible=True) # Word-Datei File Component anzeigen
                    else:
                        return gr.File(label="Fehler beim Erstellen der Word-Datei", visible=False) # Fehlermeldung anzeigen
                else:
                    return gr.File(label="Diskussions-ID fehlt. Starten Sie zuerst eine Konversation.", visible=False) # Fehlermeldung anzeigen

            def upvote_event(r_state: dict):
                did = r_state.get("discussion_id")
                itn = r_state.get("iteration")
                agn = r_state.get("agent_name")
                if did and itn and agn:
                    self.rate_agent_response(did, itn, agn, "upvote")
                    return "👍 Upvote gegeben"
                return "Fehler beim Upvote (fehlende Daten)."

            def downvote_event(r_state: dict):
                did = r_state.get("discussion_id")
                itn = r_state.get("iteration")
                agn = r_state.get("agent_name")
                if did and itn and agn:
                    self.rate_agent_response(did, itn, agn, "downvote")
                    return "👎 Downvote gegeben"
                return "Fehler beim Downvote (fehlende Daten)."

            # Klick auf "Konversation starten"
            start_btn.click(
                fn=process_input_custom_agents,
                inputs=[
                    topic_input,
                    chatbot,
                    rating_info,
                    user_state,
                    *agent_checkboxes,
                    *agent_personalities_radios,
                    iteration_slider,
                    level_radio,
                    lang_radio,
                    gr.State(None)
                ],
                outputs=[
                    chatbot,
                    formatted_output_md,
                    gr.State(),
                    rating_row,
                    rating_info,
                    word_file_output # Word-Datei Output Komponente hinzufügen
                ]
            )

            # Klick auf "Diskussion speichern"
            save_btn.click(
                fn=save_discussion_manually,
                inputs=[
                    chatbot,
                    topic_input,
                    gr.State(None),
                    formatted_output_md,
                    gr.State(),
                    user_state
                ],
                outputs=[register_status_label]
            )

            # Klick auf "Chat als Word speichern"
            word_save_btn.click(
                fn=save_chat_word_event,
                inputs=[chatbot, gr.State()], # Benötigt Chatbot-History und Discussion-ID
                outputs=[word_file_output] # Word-Datei Output Komponente
            )


            upvote_btn.click(upvote_event, inputs=[rating_info], outputs=[rating_label])
            downvote_btn.click(downvote_event, inputs=[rating_info], outputs=[rating_label])


class AgentConvoTabStandalone: # Standalone-Klasse für Testzwecke
    def run(self):
        """Startet die Agenten-Konversationsschnittstelle als eigenständige Gradio-Anwendung."""
        agent_convo_tab = AgentConversationTab() # Instanziere AgentConversationTab
        demo = gr.Blocks(title="Agenten-Konversation - Standalone - CipherCore")
        with demo:
            agent_convo_tab.build_tab() # Verwende die build_tab-Methode der Instanz
        demo.launch(share=True)


if __name__ == "__main__":
    standalone_tab = AgentConvoTabStandalone() # Verwende die Standalone-Klasse zum Testen
    standalone_tab.run()